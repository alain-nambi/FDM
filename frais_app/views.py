from django.shortcuts import render, redirect
from django.views.generic import ListView, FormView, TemplateView
from django.urls import reverse_lazy
from .models import Mission, Expense, Technician, Worker, MissionFile
from django.views import View
from decimal import Decimal
from django.http import JsonResponse
from django.contrib.auth.views import LoginView, LogoutView
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q
from django.db.models.functions import TruncMonth, ExtractMonth, ExtractYear
from django.utils.dateparse import parse_date
from django.contrib import messages
from django.contrib.auth.models import User
from datetime import datetime
import calendar
import re
from django.shortcuts import get_object_or_404
from django.urls import reverse
from django.core.exceptions import PermissionDenied
from django.template.loader import render_to_string
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from weasyprint import HTML
from django.http import HttpResponse
from io import BytesIO,StringIO
import tempfile
from datetime import date
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl import Workbook
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from django.core.mail import send_mail
from django.views.generic.edit import UpdateView
from django.conf import settings
import threading
from django.core.mail import send_mail
from django.http import HttpResponseRedirect
from django.db import transaction
import io
import pandas as pd
from django.views.generic import TemplateView
from django.contrib.auth import logout


#barre de recherche reutilisable
class MissionSearchUtils:
    @staticmethod
    def filter_missions(queryset, search_query):
        mois_fr = {
            'janvier': 1, 'février': 2, 'mars': 3, 'avril': 4, 'mai': 5, 'juin': 6,
            'juillet': 7, 'août': 8, 'septembre': 9, 'octobre': 10, 'novembre': 11, 'décembre': 12
        }
        search_lower = search_query.lower()
        month_number = None
        for month_name, month_num in mois_fr.items():
            if month_name.startswith(search_lower):
                month_number = month_num
                break

        is_year = search_query.isdigit() and len(search_query) == 4
        if month_number and is_year:
            queryset = queryset.filter(
                Q(start_date__month=month_number, start_date__year=search_query) |
                Q(end_date__month=month_number, end_date__year=search_query)
            )
        elif month_number:
            queryset = queryset.filter(
                Q(start_date__month=month_number) |
                Q(end_date__month=month_number)
            )
        elif is_year:
            queryset = queryset.filter(
                Q(start_date__year=search_query) |
                Q(end_date__year=search_query)
            )
        else:
            search_terms = search_query.split(' ')
            facturation_search = None
            if search_query.lower() in ['facturé', 'facture', 'facturée', 'oui', 'yes']:
                facturation_search = True
            elif search_query.lower() in ['non facturé', 'non facture', 'non facturée', 'non', 'no']:
                facturation_search = False

            if facturation_search is not None:
                queryset = queryset.filter(facturation=facturation_search)
            elif len(search_terms) == 1:
                queryset = queryset.filter(
                    Q(id__icontains=search_query) |
                    Q(mission_details__icontains=search_query) |
                    Q(location__icontains=search_query) |
                    Q(techniciens__first_name__icontains=search_query) |
                    Q(techniciens__last_name__icontains=search_query)
                ).distinct()
            else:
                queryset = queryset.filter(
                    Q(id__icontains=search_query) |
                    Q(location__icontains=search_query) |
                    Q(mission_details__icontains=search_query) |
                    (
                        (Q(techniciens__first_name__icontains=search_terms[0]) & Q(techniciens__last_name__icontains=search_terms[1])) |
                        (Q(techniciens__first_name__icontains=search_terms[1]) & Q(techniciens__last_name__icontains=search_terms[0]))
                    )
                ).distinct()
        return queryset


#Pagination reutilisable    
class PaginationUtils:
    @staticmethod
    def paginate_queryset(queryset, request, per_page_default=10):
        """
        Paginate a queryset based on the request parameters.

        :param queryset: Queryset à paginer
        :param request: Objet request contenant les paramètres GET
        :param per_page_default: Nombre d'éléments par page par défaut
        :return: Un objet paginé
        """
        per_page = request.GET.get('per_page', per_page_default)
        try:
            per_page = int(per_page)
        except (ValueError, TypeError):
            per_page = per_page_default

        paginator = Paginator(queryset, per_page)
        page = request.GET.get('page', 1)
        try:
            paginated_queryset = paginator.page(page)
        except PageNotAnInteger:
            paginated_queryset = paginator.page(1)
        except EmptyPage:
            paginated_queryset = paginator.page(paginator.num_pages)

        return paginated_queryset   
    
#liste des missions

class MissionListView(LoginRequiredMixin, View):
    def get(self, request, *args, **kwargs):
        all_missions = Mission.objects.exclude(status__in=['VALIDATED', 'CLOSED']).order_by('-id')
        search_query = request.GET.get('search', '')
        if search_query:
            all_missions = MissionSearchUtils.filter_missions(all_missions, search_query)

        # Utilisation de PaginationUtils
        missions = PaginationUtils.paginate_queryset(all_missions, request)

        # Récupère les techniciens pour le formulaire
        technicians = Technician.objects.values('id', 'first_name', 'last_name')
        # Compter les missions avec statut NEW
        new_missions = Mission.objects.filter(status='NEW').count()
        validated_missions = Mission.objects.filter(status='VALIDATED').count()
        refused_missions = Mission.objects.filter(status='REFUSED').count()
        total_missions = Mission.objects.all().count()
        
        context = {
            'missions': missions,
            'technicians': technicians,
            'active_tab': 'missions',  # Pour le style lorsqu'on clique sur historique ou accueil
            'new_missions': new_missions,
            'validated_missions': validated_missions,
            'refused_missions': refused_missions,
            'total_missions': total_missions,
        }
        return render(request, 'index.html', context)
        
    # stockage des données de la mission dans la base
    def post(self, request, *args, **kwargs):
        bluedesk_link = request.POST.get('bluedesk_link')
        mission_details = request.POST.get('mission_details')
        start_date = request.POST.get('start_date')
        start_hour = request.POST.get('start_hour')
        end_date = request.POST.get('end_date')
        end_hour = request.POST.get('end_hour')
        location = request.POST.get('location')
        facturation = request.POST.get('facturation') == 'on'
        hosting_days = int(request.POST.get('hosting_days', 0))
        overnight_rate = Decimal(request.POST.get('overnight_rate', 0))
        meal_costs = Decimal(request.POST.get('meal_costs', 0))
        transport = request.POST.get('transport')
        shipping_costs = Decimal(request.POST.get('shipping_costs', 0))
        various_expenses_details = request.POST.get('various_expenses_details')
        various_expenses_price = Decimal(request.POST.get('various_expenses_price', 0))
        
        # Créer une nouvelle mission avec les données récupérées
        mission = Mission.objects.create(
            bluedesk_link=bluedesk_link,
            mission_details=mission_details,
            start_date=start_date,
            start_hour=start_hour,
            end_date=end_date,
            end_hour=end_hour,
            location=location,
            facturation=facturation,
            created_by=request.user,
            updated_by=request.user  # Ajouter l'utilisateur qui crée comme le premier à mettre à jour
        )
        
        def send_email_async(subject, message, from_email, recipient_list):
            threading.Thread(
                 target=send_mail,
                 args=(subject, message, from_email, recipient_list),
                 kwargs={'fail_silently': False}
             ).start()
        
        techniciens_ids = request.POST.getlist('techniciens')
        for tech_id in techniciens_ids:
            technician = Technician.objects.get(id=tech_id)
            mission.techniciens.add(technician)
            
        # Créer une nouvelle dépense associée à la mission
        Expense.objects.create(
            mission=mission,
            hosting_days=hosting_days,
            overnight_rate=overnight_rate,
            meal_costs=meal_costs,
            transport=transport,
            shipping_costs=shipping_costs,
            various_expenses_details=various_expenses_details,
            various_expenses_price=various_expenses_price
        )
        
        # Envoyer une notification par e-mail de manière asynchrone
        subject = "Nouvelle mission créée"
        message = f"Une nouvelle mission a été créée par {request.user.username}.\n\nDétails : {mission.mission_details}"
        recipient_list = ['mihajarazafimahazoson@gmail.com']  # Remplacez par l'adresse e-mail souhaitée à qui envoyer le mail, normalement DG
        send_email_async(subject, message, 'mihaja356@gmail.com', recipient_list)
        
        # Ajouter un message de succès
        messages.success(request, "La mission a été créée avec succès!")
        
        return redirect('missions')


# historiques des missions validés 
class HistoryView(TemplateView):
    template_name = 'history.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Récupération des missions validées ET clôturées pour l'affichage du tableau
        history_missions = Mission.objects.filter(
            status__in=['VALIDATED', 'CLOSED']
        ).order_by('-id')
        
        # Recherche
        search_query = self.request.GET.get('search', '')
        if search_query:
            history_missions = MissionSearchUtils.filter_missions(history_missions, search_query)
            
        # Utilisation de PaginationUtils
        missions = PaginationUtils.paginate_queryset(history_missions, self.request)
        
        # Ajout des compteurs pour account.html
        context['new_missions'] = Mission.objects.filter(status='NEW').count()
        context['validated_missions'] = Mission.objects.filter(status='VALIDATED').count()
        context['refused_missions'] = Mission.objects.filter(status='REFUSED').count()
        context['closed_missions'] = Mission.objects.filter(status='CLOSED').count()
        context['total_missions'] = Mission.objects.all().count()
        context['missions'] = missions
        context['active_tab'] = 'history'
        
        return context
    
    def post(self, request, *args, **kwargs):
        """Méthode pour gérer la clôture de mission directement depuis la vue History."""
        mission_id = request.POST.get('mission_id')
        
        # Détermine l'action demandée
        action = request.POST.get('action', 'close')  # Par défaut, c'est une clôture
        
        if mission_id:
            try:
                mission = Mission.objects.get(id=mission_id)
                
                # Gérer différentes actions possibles
                if action == 'close':
                    # Clôture de mission
                    mission.status = 'CLOSED'
                    mission.save()
                    messages.success(
                        request, 
                        f'La mission #{mission.id} a été clôturée avec succès.',
                        extra_tags='cloture'  # Ajout du tag pour la notification
                    )
                elif action == 'upload_file':
                    # Traitement de l'upload de fichier
                    # (à gérer selon votre logique d'upload)
                    file = request.FILES.get('file')
                    if file:
                        # Logique pour sauvegarder le fichier
                        # Ceci est un exemple, à adapter selon votre modèle
                        MissionFile.objects.create(
                            mission=mission,
                            file=file,
                            uploaded_by=request.user
                        )
                        messages.success(
                            request, 
                            f'Le fichier a été ajouté à la mission #{mission.id} avec succès.',
                            extra_tags='fichier'  # Ajout du tag pour la notification
                        )
                    else:
                        messages.warning(
                            request, 
                            'Aucun fichier sélectionné.',
                            extra_tags='fichier'  # Ajout du tag pour la notification
                        )
                
            except Mission.DoesNotExist:
                messages.error(
                    request, 
                    f'Mission #{mission_id} introuvable.',
                    extra_tags='cloture'  # Ajout du tag pour la notification
                )
        
        return redirect('history')
    
    
# Inscription
class RegisterView(TemplateView):
    template_name = "register.html"
    success_url = reverse_lazy("login")

    def post(self, request, *args, **kwargs):
        if request.method == "POST":
            username = request.POST["username"]
            first_name = request.POST["first_name"]
            last_name = request.POST["last_name"]
            poste = request.POST["poste"]
            password = request.POST["password"]
            email = request.POST["email"]

            if User.objects.filter(username=username).exists():
                messages.error(request, "Ce nom d'utilisateur existe déjà")
            else:
                user = User.objects.create_user(
                    username=username,
                    email=email, password=password,
                    first_name=first_name,
                    last_name=last_name
                    )
                
                Worker.objects.create(user_id=user, poste=poste)
                 # Si l'utilisateur est un technicien, créer aussi une entrée dans le modèle Technician
                if poste == "Techniciens":
                    Technician.objects.create(
                        first_name=first_name,
                        last_name=last_name
                    )
                messages.success(request,"Inscription réussie, connectez-vous maintenant")
                return redirect("login")

        return self.get(request, *args, **kwargs)


# Connexion
class CustomLoginView(LoginView):
    template_name = "login.html"

    def form_valid(self, form):
        response = super().form_valid(form)
        user = self.request.user

        # Vérifier si l'utilisateur a un employé associé
        return redirect("missions")


# Déconnexion
class CustomLogoutView(LogoutView):
    def dispatch(self, request, *args, **kwargs):
        logout(request)
        messages.success(request, "Vous avez été déconnecté.")
        super().dispatch(request, *args, **kwargs)
        return redirect('login')
    
    
    
# class pour mettre à jour les données entrés par l'utilisateur 
class EditMissionView(View):
    def post(self, request, mission_id, *args, **kwargs):
        mission = Mission.objects.get(id=mission_id)
        
        # Mettre le statut à "NEW"
        mission.status = 'NEW'
        
        # Récupérer les données du formulaire
        bluedesk_link = request.POST.get('bluedesk_link')
        mission_details = request.POST.get('mission_details')
        start_date = request.POST.get('start_date')
        start_hour = request.POST.get('start_hour')
        end_date = request.POST.get('end_date')
        end_hour = request.POST.get('end_hour')
        location = request.POST.get('location')
        facturation = request.POST.get('facturation') == 'on'
        
        
        # Mettre à jour la mission
        mission.bluedesk_link = bluedesk_link
        mission.mission_details = mission_details
        mission.start_date = start_date
        mission.start_hour = start_hour
        mission.end_date = end_date
        mission.end_hour = end_hour
        mission.location = location
        mission.facturation = facturation
        mission.updated_by = request.user
        mission.save()
        
        # Mise à jour des techniciens
        mission.techniciens.clear()
        techniciens_ids = request.POST.getlist('techniciens')
        for tech_id in techniciens_ids:
            technician = Technician.objects.get(id=tech_id)
            mission.techniciens.add(technician)
        
        # Mise à jour des dépenses
        try:
            expense = mission.depenses.first()  # Suppose que mission.depenses est le related_name dans le modèle Expense
            
            hosting_days = int(request.POST.get('hosting_days', 0))
            overnight_rate = Decimal(request.POST.get('overnight_rate', 0))
            meal_costs = Decimal(request.POST.get('meal_costs', 0))
            transport = request.POST.get('transport')
            shipping_costs = Decimal(request.POST.get('shipping_costs', 0))
            various_expenses_details = request.POST.get('various_expenses_details')
            various_expenses_price = Decimal(request.POST.get('various_expenses_price', 0))
            
            # Mettre à jour les valeurs
            expense.hosting_days = hosting_days
            expense.overnight_rate = overnight_rate
            expense.meal_costs = meal_costs
            expense.transport = transport
            expense.shipping_costs = shipping_costs
            expense.various_expenses_details = various_expenses_details
            expense.various_expenses_price = various_expenses_price
            
            # Recalcul des totaux (si nécessaire)
            # Note: Ceci dépend de votre modèle et de vos calculs spécifiques
            start_dt = datetime.strptime(start_date, '%Y-%m-%d')
            end_dt = datetime.strptime(end_date, '%Y-%m-%d')
            days_diff = (end_dt - start_dt).days + 1  # +1 pour inclure le jour de fin
            
            expense.total_hosting = hosting_days * overnight_rate
            expense.total_meal_costs = meal_costs * days_diff
            expense.total_expenses = (
                expense.total_hosting + 
                expense.total_meal_costs + 
                shipping_costs + 
                various_expenses_price
            )
            
            expense.save()
        except Exception as e:
            # Gérer l'erreur si nécessaire
            pass
        
        return redirect('missions')
    
    
# Fonction utilitaire pour exécuter des tâches asynchrones
def run_async_task(target, *args, **kwargs):
    threading.Thread(target=target, args=args, kwargs=kwargs).start()
    
# Pour la validation des missions
class ValidateMissionView(View):
    def post(self, request, *args, **kwargs):
        if not request.user.has_perm('frais_app.can_validate_mission'):
            raise PermissionDenied

        mission_id = request.POST.get('mission_id')
        mission = get_object_or_404(Mission, id=mission_id)

        # Mettre à jour le statut
        mission.status = 'VALIDATED'
        mission.validated_at = datetime.now()
        mission.save()

        # Préparer l'envoi d'e-mail
        subject = f"Demande de validation frais de mission à {mission.location}"
        message = f"La mission #{mission.id} a été validée avec succès."
        notification_email = mission.created_by.email

        # Exécuter l'envoi d'e-mail de manière asynchrone
        run_async_task(
            send_mail,
            subject,
            message,
            'mihajarazafimahazoson@gmail.com',  # Adresse d'expéditeur
            [notification_email],  # Liste des destinataires
            fail_silently=True
        )

        messages.success(request, "La mission a été validée avec succès.")
        return HttpResponseRedirect(reverse('missions'))

# Pour le refus de la mission
class RefuseMissionView(View):
    def post(self, request, *args, **kwargs):
        if not request.user.has_perm('frais_app.can_refuse_mission'):
            raise PermissionDenied

        mission_id = request.POST.get('mission_id')
        refusal_reason = request.POST.get('refusal_reason', '')

        if not refusal_reason.strip():
            messages.error(request, "Veuillez saisir un motif de refus.")
            return HttpResponseRedirect(reverse('missions'))

        mission = get_object_or_404(Mission, id=mission_id)

        # Mettre à jour le statut
        mission.status = 'REFUSED'
        mission.refusal_reason = refusal_reason
        mission.refused_at = datetime.now()
        mission.save()

        # Préparer l'envoi d'e-mail
        subject = f"Demande de validation frais de mission à {mission.location}"
        message = f"La mission #{mission.id} a été refusée pour la raison suivante : {refusal_reason}."
        notification_email = mission.created_by.email

        # Exécuter l'envoi d'e-mail de manière asynchrone
        run_async_task(
            send_mail,
            subject,
            message,
            'mihajarazafimahazoson@gmail.com',  # Adresse d'expéditeur
            [notification_email],  # Liste des destinataires
            fail_silently=True
        )

        messages.success(request, "La mission a été refusée avec succès.")
        return HttpResponseRedirect(reverse('missions'))
    
  
#class pour le telechargement du pdf dans le modal details 
class GeneratePDFView(View):
    def get(self, request, mission_id, *args, **kwargs):
        # Récupérer la mission
        mission = get_object_or_404(Mission, id=mission_id)
        
        # Charger le template HTML
        context = {
            'mission': mission,
            'expenses': mission.depenses.all() 
        }
        
        html_string = render_to_string('pdf_template.html', context)
        
        # Créer un fichier temporaire pour stocker le PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            temp_filename = tmp.name
        
        # Générer le PDF avec WeasyPrint
        html = HTML(string=html_string)
        html.write_pdf(temp_filename)
        
        # Lire le fichier temporaire et le renvoyer dans la réponse
        with open(temp_filename, 'rb') as f:
            pdf_content = f.read()
            
        # Retourner le fichier PDF en tant que réponse HTTP
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="mission_{}.pdf"'.format(mission.id)
        
        
        return response
    

#class pour le telechargement du pdf de toutes les missions
class ExportMissionsPDFView(View):
    def get(self, request):
        # Récupérer toutes les missions
        missions = Mission.objects.all().prefetch_related('depenses', 'techniciens')
        
        # Préparer les données pour le résumé
        validated_count = missions.filter(status='VALIDATED').count()
        new_count = missions.filter(status='NEW').count()
        refused_count = missions.filter(status='REFUSED').count()
        
        # Calculer le total des dépenses
        total_expenses = 0
        for mission in missions:
            for expense in mission.depenses.all():
                total_expenses += expense.total_expenses
        
        # Préparer le contexte
        context = {
            'missions': missions,
            'validated_count': validated_count,
            'new_count': new_count,
            'refused_count': refused_count,
            'total_expenses': total_expenses,
            'today': date.today(),
        }
        
        # Rendre le HTML
        html_string = render_to_string('missions_pdf_export.html', context)
        
        # Générer le PDF
        html = HTML(string=html_string)
        pdf_file = html.write_pdf()
        
        # Créer la réponse HTTP
        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="missions_export.pdf"'
        
        return response
    
    
#class pour l'export Excel d'une mission 
class ExportMissionExcelView(View):
    def get(self, request, mission_id):
        # Récupérer la mission spécifique
        try:
            mission = Mission.objects.prefetch_related('depenses', 'techniciens').get(id=mission_id)
        except Mission.DoesNotExist:
            return HttpResponse("Mission non trouvée", status=404)
            
        # Créer un classeur
        wb = Workbook()
        ws = wb.active
        ws.title = "Mission {}".format(mission_id)
        
        # Styles de base
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
        header_alignment = Alignment(horizontal='center')
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'), 
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        
        # Première feuille: Détails de la mission
        headers = ['ID', 'Détails', 'Techniciens', 'Lieu', 'Date de début',
                  'Date de fin', 'Statut', 'Total des dépenses(Ar)']
                  
        # Appliquer les en-têtes
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
            ws.column_dimensions[chr(64 + col_num)].width = 20
        
        # Données de la mission
        tech_list = ', '.join(["{} {}".format(tech.first_name, tech.last_name) for tech in mission.techniciens.all()])
        total_expenses = sum(expense.total_expenses for expense in mission.depenses.all())
        
        # Mapper le statut
        status_mapping = {
            'NEW': 'Nouvelle',
            'VALIDATED': 'Validée',
            'REFUSED': 'Refusée'
        }
        status_display = status_mapping.get(mission.status, mission.status)
        
        # Ajouter les données
        row = [
            mission.id,
            mission.mission_details[:100],
            tech_list,
            mission.location,
            mission.start_date.strftime('%d/%m/%Y'),
            mission.end_date.strftime('%d/%m/%Y'),
            status_display,
            "{:.2f}".format(total_expenses)
        ]
        
        for col_num, cell_value in enumerate(row, 1):
            cell = ws.cell(row=2, column=col_num)
            cell.value = cell_value
            cell.border = border
            
            # Style du statut
            if col_num == 7:  # Colonne statut
                if cell_value == 'Validée':
                    cell.font = Font(color='27AE60', bold=True)
                elif cell_value == 'Refusée':
                    cell.font = Font(color='E74C3C', bold=True)
                elif cell_value == 'Nouvelle':
                    cell.font = Font(color='F39C12', bold=True)
        
        # Deuxième feuille: Dépenses
        ws_expenses = wb.create_sheet(title="Dépenses")
        
        # En-têtes pour les dépenses
        expense_headers = [
            'ID', 'Jours d\'hébergement', 'Prix nuitée', 'Total hébergement', 
            'Coût repas', 'Total repas', 'Transport', 'Frais de transport', 
            'Divers', 'Frais divers', 'Total'
        ]
        
        # Appliquer les en-têtes
        for col_num, header in enumerate(expense_headers, 1):
            cell = ws_expenses.cell(row=1, column=col_num)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
            ws_expenses.column_dimensions[chr(64 + col_num)].width = 18
        
        # Ajouter les dépenses
        for row_num, expense in enumerate(mission.depenses.all(), 2):
            expense_row = [
                expense.id,
                expense.hosting_days,
                "{:.2f}".format(expense.overnight_rate),
                "{:.2f}".format(expense.total_hosting),
                "{:.2f}".format(expense.meal_costs),
                "{:.2f}".format(expense.total_meal_costs),
                expense.transport,
                "{:.2f}".format(expense.shipping_costs),
                expense.various_expenses_details,
                "{:.2f}".format(expense.various_expenses_price),
                "{:.2f}".format(expense.total_expenses)
                            ]
            
            for col_num, cell_value in enumerate(expense_row, 1):
                cell = ws_expenses.cell(row=row_num, column=col_num)
                cell.value = cell_value
                cell.border = border
        
        # Troisième feuille: Résumé
        ws_summary = wb.create_sheet(title="Résumé")
        
        ws_summary.cell(row=1, column=1).value = "Résumé de la Mission #{}".format(mission_id)
        ws_summary.cell(row=1, column=1).font = Font(bold=True, size=14)
        ws_summary.cell(row=2, column=1).value = "Généré le {}".format(datetime.now().strftime('%d/%m/%Y à %H:%M'))
        
        # Informations résumées
        summary_data = [
            ("Détails", mission.mission_details[:100]),
            ("Lieu", mission.location),
            ("Période", "{} - {}".format(mission.start_date.strftime('%d/%m/%Y'), mission.end_date.strftime('%d/%m/%Y'))),
            ("Statut", status_display),
            ("Techniciens", tech_list),
            ("Nombre de dépenses", mission.depenses.count()),
            ("Total des dépenses", "{:.2f} Ar".format(total_expenses))
        ]
        
        for row_num, (label, value) in enumerate(summary_data, 4):
            ws_summary.cell(row=row_num, column=1).value = label
            ws_summary.cell(row=row_num, column=2).value = value
            if label == "Statut":
                if value == 'Validée':
                    ws_summary.cell(row=row_num, column=2).font = Font(color='27AE60', bold=True)
                elif value == 'Refusée':
                    ws_summary.cell(row=row_num, column=2).font = Font(color='E74C3C', bold=True)
                elif value == 'Nouvelle':
                    ws_summary.cell(row=row_num, column=2).font = Font(color='F39C12', bold=True)
        
        # Ajuster largeur des colonnes
        for col in range(1, 3):
            ws_summary.column_dimensions[chr(64 + col)].width = 25
        
        # Sauvegarder dans un buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        # Créer la réponse HTTP
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=mission_{}_export.xlsx'.format(mission_id)
        
        return response


#class pour l'export CSV 
class ExportMissionCSVView(View):
    def get(self, request, mission_id):
        try:
            mission = Mission.objects.prefetch_related('depenses', 'techniciens').get(id=mission_id)
        except Mission.DoesNotExist:
            return HttpResponse("Mission non trouvée", status=404)

        # Configuration du CSV
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename=mission_{}_export_{}.csv'.format(
    mission_id, datetime.now().strftime("%Y%m%d_%H%M")
)
        
        # Création du writer avec des options améliorées
        writer = csv.writer(response, delimiter=';')
        
        # Écriture de l'en-tête principal
        writer.writerow(['EXPORT DE MISSION - DÉTAILS'])
        writer.writerow([])
        
        # Section Informations de base
        writer.writerow(['INFORMATIONS GÉNÉRALES'])
        writer.writerow([
            'ID Mission', 'Détails', 'Lieu', 
            'Date début', 'Date fin', 'Statut', 'Total dépenses (Ar)'
        ])
        
        status_mapping = {
            'NEW': 'Nouvelle',
            'VALIDATED': 'Validée',
            'REFUSED': 'Refusée'
        }
        
        total_expenses = sum(expense.total_expenses for expense in mission.depenses.all())
        tech_list = ' | '.join(["{} {}".format(tech.first_name, tech.last_name) for tech in mission.techniciens.all()])
        
        writer.writerow([
            mission.id,
            mission.mission_details,
            mission.location,
            mission.start_date.strftime('%d/%m/%Y'),
            mission.end_date.strftime('%d/%m/%Y'),
            status_mapping.get(mission.status, mission.status),
            "{:,.2f}".format(total_expenses).replace(',', ' ')
        ])
        
        writer.writerow([])
        
        # Section Techniciens
        writer.writerow(['TECHNICIENS ASSIGNÉS'])
        writer.writerow(['Liste des techniciens'])
        writer.writerow([tech_list])
        writer.writerow([])
        
        # Section Dépenses (avec en-tête détaillé)
        writer.writerow(['DÉTAIL DES DÉPENSES'])
        headers = [
            'ID', 'Jours hébergement', 'Prix/nuit (Ar)', 'Total hébergement (Ar)',
            'Coût repas/jour (Ar)', 'Total repas (Ar)', 'Type transport',
            'Frais transport (Ar)', 'Détails divers', 'Frais divers (Ar)', 'Total (Ar)'
        ]
        writer.writerow(headers)
        
        for expense in mission.depenses.all():
            writer.writerow([
                expense.id,
                expense.hosting_days,
                "{:,.2f}".format(expense.overnight_rate).replace(',', ' '),
                "{:,.2f}".format(expense.total_hosting).replace(',', ' '),
                "{:,.2f}".format(expense.meal_costs).replace(',', ' '),
                "{:,.2f}".format(expense.total_meal_costs).replace(',', ' '),
                expense.transport,
                "{:,.2f}".format(expense.shipping_costs).replace(',', ' '),
                expense.various_expenses_details or 'N/A',
                "{:,.2f}".format(expense.various_expenses_price).replace(',', ' '),
                "{:,.2f}".format(expense.total_expenses).replace(',', ' ')
            ])
        
        # Section Totaux
        writer.writerow([])
        writer.writerow(['TOTAUX'])
        writer.writerow(['Total général des dépenses:', "{:,.2f} Ar".format(total_expenses).replace(',', ' ')])
        
        # Pied de page
        writer.writerow([])
        writer.writerow(['Export généré le', datetime.now().strftime('%d/%m/%Y à %H:%M')])
        
        return response


#class pour l'export Word
class ExportMissionDocxView(View):
    def get(self, request, mission_id):
        try:
            # Récupérer la mission spécifique
            mission = Mission.objects.prefetch_related('depenses', 'techniciens').get(id=mission_id)
        except Mission.DoesNotExist:
            return HttpResponse("Mission non trouvée", status=404)
            
        # Créer un nouveau document
        doc = Document()
        
        # Configurer le style du document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Ajouter un titre
        title = doc.add_heading('Rapport de Mission #{}'.format(mission_id), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter la date de génération
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run('Généré le {}'.format(datetime.now().strftime("%d/%m/%Y à %H:%M")))
        
        # Mapper les statuts pour l'affichage
        status_mapping = {
            'NEW': 'Nouvelle',
            'VALIDATED': 'Validée',
            'REFUSED': 'Refusée'
        }
        
        # Formater les dates
        start_date = mission.start_date.strftime('%d/%m/%Y')
        end_date = mission.end_date.strftime('%d/%m/%Y')
        
        # Obtenir la liste des techniciens
        tech_list = ', '.join(["{} {}".format(tech.first_name, tech.last_name) for tech in mission.techniciens.all()])
        
        # Total des dépenses
        total_mission_expenses = sum(expense.total_expenses for expense in mission.depenses.all())
        
        # Récupérer l'affichage du statut
        status_display = status_mapping.get(mission.status, mission.status)
        
        # Ajouter les détails de la mission
        doc.add_heading('Détails de la Mission', level=1)
        
        details_table = doc.add_table(rows=6, cols=2)
        details_table.style = 'Light Grid Accent 1'
        
        # Remplir le tableau des détails
        details_rows = [
            ('Détails:', mission.mission_details),
            ('Lieu:', mission.location),
            ('Date de début:', start_date),
            ('Date de fin:', end_date),
            ('Techniciens:', tech_list),
            ('Statut:', status_display)
        ]
        
        for i, (label, value) in enumerate(details_rows):
            cell = details_table.cell(i, 0)
            cell.text = label
            cell.width = Inches(1.5)
            cell.paragraphs[0].runs[0].bold = True
            
            cell = details_table.cell(i, 1)
            cell.text = value
            
            # Colorer le statut
            if label == 'Statut:':
                run = cell.paragraphs[0].runs[0]
                if value == 'Validée':
                    run.font.color.rgb = RGBColor(39, 174, 96)  # Vert
                    run.bold = True
                elif value == 'Refusée':
                    run.font.color.rgb = RGBColor(231, 76, 60)  # Rouge
                    run.bold = True
                elif value == 'Nouvelle':
                    run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
                    run.bold = True
        
        # Ajouter des dépenses si elles existent
        expenses = mission.depenses.all()
        if expenses:
            doc.add_heading('Dépenses', level=2)
            
            # Créer un tableau pour les dépenses
            expense_table = doc.add_table(rows=1, cols=6)
            expense_table.style = 'Table Grid'
            
            # En-têtes
            headers = ['Nuitées', 'Hébergement', 'Repas', 'Transport', 'Divers', 'Total']
            for i, header in enumerate(headers):
                cell = expense_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ajouter chaque dépense
            for expense in expenses:
                row_cells = expense_table.add_row().cells
                row_cells[0].text = "{} jours × {:.2f} Ar".format(expense.hosting_days, expense.overnight_rate)
                row_cells[1].text = "{:.2f} Ar".format(expense.total_hosting)
                row_cells[2].text = "{:.2f} Ar".format(expense.total_meal_costs)
                row_cells[3].text = "{:.2f} Ar ({})".format(expense.shipping_costs, expense.transport)
                row_cells[4].text = "{:.2f} Ar".format(expense.various_expenses_price)
                row_cells[5].text = "{:.2f} Ar".format(expense.total_expenses)
                                
                # Aligner à droite les cellules
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Ajouter une ligne pour le total de la mission
            row_cells = expense_table.add_row().cells
            row_cells[4].text = "TOTAL:"
            row_cells[4].paragraphs[0].runs[0].bold = True
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            row_cells[5].text = "{:.2f} Ar".format(total_mission_expenses)
            row_cells[5].paragraphs[0].runs[0].bold = True
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Ajouter les détails des dépenses diverses si présents
            has_details = any(expense.various_expenses_details for expense in expenses)
            if has_details:
                doc.add_heading('Détails des frais divers', level=3)
                for expense in expenses:
                    if expense.various_expenses_details:
                        p = doc.add_paragraph()
                        p.add_run("• {}: ".format(expense.various_expenses_details)).bold = True
                        p.add_run("{:.2f} Ar".format(expense.various_expenses_price))
        
        # Sauvegarder dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Créer la réponse HTTP
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=mission_{}_rapport.docx'.format(mission_id)
        
        return response
    
    
    
    
#class pour l'enregistrement des fichiers 
class UploadMissionFileView(View):
    def post(self, request, mission_id):
        mission = get_object_or_404(Mission, id=mission_id)
        files = request.FILES.getlist('files')
        description = request.POST.get('description', 'Pas de description')
        
        if files:
            for file in files:
                MissionFile.objects.create(
                    mission=mission,
                    file=file,
                    file_description=description,
                    uploaded_by=request.user  # Enregistre l'utilisateur qui a uploadé
                )
            messages.success(request, "{} fichier(s) ajouté(s) à la mission #{}".format(len(files), mission_id))
        else:
            messages.error(request, "Aucun fichier n'a été sélectionné")
            
        # Rediriger vers la page d'où provient la requête
        return redirect('history')



#pour la création d'un technicien
class CreateTechnicianView(LoginRequiredMixin, TemplateView):
    template_name = 'create_technician.html'
    
    def get(self, request, *args, **kwargs):
        context = {}
        
        # Récupérer les messages de session et les ajouter au contexte
        if 'single_success' in request.session:
            context['single_success'] = request.session.pop('single_success')
        
        if 'form_errors' in request.session:
            context['form_errors'] = request.session.pop('form_errors')
            
        if 'file_success' in request.session:
            context['file_success'] = request.session.pop('file_success')
            
        if 'file_errors' in request.session:
            context['file_errors'] = request.session.pop('file_errors')
            
        return render(request, self.template_name, context)
    
    def post(self, request, *args, **kwargs):
        form_type = request.POST.get('form_type')
        
        if form_type == 'single':
            # Traitement de l'ajout d'un seul technicien
            return self._handle_single_technician(request)
        elif form_type == 'file':
            # Traitement de l'import par fichier
            return self._handle_file_import(request)
        
        # Si form_type n'est pas reconnu, rediriger vers la page
        return redirect('create_technician')
    
    def _handle_single_technician(self, request):
        """Gère l'ajout d'un technicien unique"""
        last_name = request.POST.get('last_name')
        first_name = request.POST.get('first_name')
        matricule = request.POST.get('matricule')
        
        # Validation des données
        errors = []
        if not last_name:
            errors.append("Le nom est requis")
        if not first_name:
            errors.append("Le prénom est requis")
        
        if errors:
            request.session['form_errors'] = errors
            return redirect('create_technician')
        
        # Création du technicien
        try:
            Technician.objects.create(
                last_name=last_name,
                first_name=first_name,
                matricule=matricule if matricule else None
            )
            request.session['single_success'] = True
        except Exception as e:
            errors.append(f"Erreur lors de l'enregistrement: {str(e)}")
            request.session['form_errors'] = errors
        
        return redirect('create_technician')
    
    def _handle_file_import(self, request):
        """Gère l'import de techniciens depuis un fichier"""
        if 'file' not in request.FILES:
            request.session['file_errors'] = ["Aucun fichier n'a été soumis"]
            return redirect('create_technician')
        
        file = request.FILES['file']
        
        # Vérifier l'extension du fichier
        if not file.name.endswith(('.xlsx', '.xls', '.csv')):
            request.session['file_errors'] = ["Format de fichier non pris en charge. Utilisez XLSX, XLS ou CSV."]
            return redirect('create_technician')
        
        try:
            # Lecture du fichier
            if file.name.endswith('.csv'):
                df = pd.read_csv(file)
            else:
                df = pd.read_excel(file)
            
            # Vérification des colonnes requises
            required_columns = ['Nom', 'Prénom']
            if not all(col in df.columns for col in required_columns):
                request.session['file_errors'] = ["Le fichier doit contenir les colonnes 'Nom' et 'Prénom'"]
                return redirect('create_technician')
            
            # Import des techniciens en une seule transaction
            technicians_added = self._import_technicians_from_dataframe(df)
            request.session['file_success'] = f"{technicians_added} technicien(s) ont été importés avec succès."
            
        except Exception as e:
            request.session['file_errors'] = [f"Erreur lors de l'importation du fichier: {str(e)}"]
        
        return redirect('create_technician')
    
    @transaction.atomic
    def _import_technicians_from_dataframe(self, df):
        """Importe les techniciens depuis un DataFrame pandas"""
        count = 0
        for _, row in df.iterrows():
            # Extraire les données du DataFrame
            last_name = row.get('Nom', '').strip()
            first_name = row.get('Prénom', '').strip()
            
            # Vérifier les valeurs obligatoires
            if not last_name or not first_name:
                continue
            
            # Gérer le matricule s'il existe
            matricule = None
            if 'Matricule' in row and pd.notna(row['Matricule']):
                matricule = str(row['Matricule']).strip()
            
            # Créer le technicien
            Technician.objects.create(
                last_name=last_name,
                first_name=first_name,
                matricule=matricule
            )
            count += 1