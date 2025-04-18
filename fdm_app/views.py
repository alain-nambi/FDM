from django.shortcuts import render, redirect
from django.views.generic import ListView, FormView, TemplateView
from django.urls import reverse_lazy
from .models import Mission, Expense, Technician, Worker
from django.views import View
from decimal import Decimal
from django.http import JsonResponse
from django.contrib.auth.views import LoginView, LogoutView
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q
from django.db.models.functions import TruncMonth, ExtractMonth, ExtractYear
from django.utils.dateparse import parse_date
from django.contrib import messages
from django.contrib.auth.models import User
from datetime import datetime
import calendar
import re
from django.shortcuts import get_object_or_404
from django.urls import reverse
from django.core.exceptions import PermissionDenied
from django.template.loader import render_to_string
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from weasyprint import HTML
from django.http import HttpResponse
from io import BytesIO,StringIO
import tempfile
from datetime import date
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl import Workbook
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from django.core.mail import send_mail
from django.views.generic.edit import UpdateView
from django.conf import settings


#barre de recherche reutilisable
class MissionSearchUtils:
    @staticmethod
    def filter_missions(queryset, search_query):
        mois_fr = {
            'janvier': 1, 'février': 2, 'mars': 3, 'avril': 4, 'mai': 5, 'juin': 6,
            'juillet': 7, 'août': 8, 'septembre': 9, 'octobre': 10, 'novembre': 11, 'décembre': 12
        }
        search_lower = search_query.lower()
        month_number = None
        for month_name, month_num in mois_fr.items():
            if month_name.startswith(search_lower):
                month_number = month_num
                break

        is_year = search_query.isdigit() and len(search_query) == 4
        if month_number and is_year:
            queryset = queryset.filter(
                Q(start_date__month=month_number, start_date__year=search_query) |
                Q(end_date__month=month_number, end_date__year=search_query)
            )
        elif month_number:
            queryset = queryset.filter(
                Q(start_date__month=month_number) |
                Q(end_date__month=month_number)
            )
        elif is_year:
            queryset = queryset.filter(
                Q(start_date__year=search_query) |
                Q(end_date__year=search_query)
            )
        else:
            search_terms = search_query.split(' ')
            facturation_search = None
            if search_query.lower() in ['facturé', 'facture', 'facturée', 'oui', 'yes']:
                facturation_search = True
            elif search_query.lower() in ['non facturé', 'non facture', 'non facturée', 'non', 'no']:
                facturation_search = False

            if facturation_search is not None:
                queryset = queryset.filter(facturation=facturation_search)
            elif len(search_terms) == 1:
                queryset = queryset.filter(
                    Q(id__icontains=search_query) |
                    Q(mission_details__icontains=search_query) |
                    Q(location__icontains=search_query) |
                    Q(techniciens__first_name__icontains=search_query) |
                    Q(techniciens__last_name__icontains=search_query)
                ).distinct()
            else:
                queryset = queryset.filter(
                    Q(id__icontains=search_query) |
                    Q(location__icontains=search_query) |
                    Q(mission_details__icontains=search_query) |
                    (
                        (Q(techniciens__first_name__icontains=search_terms[0]) & Q(techniciens__last_name__icontains=search_terms[1])) |
                        (Q(techniciens__first_name__icontains=search_terms[1]) & Q(techniciens__last_name__icontains=search_terms[0]))
                    )
                ).distinct()
        return queryset


#Pagination reutilisable    
class PaginationUtils:
    @staticmethod
    def paginate_queryset(queryset, request, per_page_default=10):
        """
        Paginate a queryset based on the request parameters.

        :param queryset: Queryset à paginer
        :param request: Objet request contenant les paramètres GET
        :param per_page_default: Nombre d'éléments par page par défaut
        :return: Un objet paginé
        """
        per_page = request.GET.get('per_page', per_page_default)
        try:
            per_page = int(per_page)
        except (ValueError, TypeError):
            per_page = per_page_default

        paginator = Paginator(queryset, per_page)
        page = request.GET.get('page', 1)
        try:
            paginated_queryset = paginator.page(page)
        except PageNotAnInteger:
            paginated_queryset = paginator.page(1)
        except EmptyPage:
            paginated_queryset = paginator.page(paginator.num_pages)

        return paginated_queryset   
    
#liste des missions
class MissionListView(View):
    def get(self, request, *args, **kwargs):
        all_missions = Mission.objects.exclude(status='VALIDATED').order_by('-id')
        search_query = request.GET.get('search', '')
        if search_query:
            all_missions = MissionSearchUtils.filter_missions(all_missions, search_query)

        # Utilisation de PaginationUtils
        missions = PaginationUtils.paginate_queryset(all_missions, request)

        # Récupère les techniciens pour le formulaire
        technicians = Technician.objects.values('id', 'first_name', 'last_name')
        # Compter les missions avec statut NEW
        new_missions = Mission.objects.filter(status='NEW').count()
        
        context = {
            'missions': missions,
            'technicians': technicians,
            'active_tab': 'missions',  # Pour le style lorsqu'on clique sur historique ou accueil
            'new_missions': new_missions 
        }
        return render(request, 'index.html', context)
        
        
    # stockage des données de la mission dans la base
    def post(self, request, *args, **kwargs):
        bluedesk_link = request.POST.get('bluedesk_link')
        mission_details = request.POST.get('mission_details')
        start_date = request.POST.get('start_date')
        start_hour = request.POST.get('start_hour')
        end_date = request.POST.get('end_date')
        end_hour = request.POST.get('end_hour')
        location = request.POST.get('location')
        facturation = request.POST.get('facturation') == 'on'
        hosting_days = int(request.POST.get('hosting_days', 0))
        overnight_rate = Decimal(request.POST.get('overnight_rate', 0))
        meal_costs = Decimal(request.POST.get('meal_costs', 0))
        transport = request.POST.get('transport')
        shipping_costs = Decimal(request.POST.get('shipping_costs', 0))
        various_expenses_details = request.POST.get('various_expenses_details')
        various_expenses_price = Decimal(request.POST.get('various_expenses_price', 0))
        
        # Créer une nouvelle mission avec les données récupérées
        mission = Mission.objects.create(
            bluedesk_link=bluedesk_link,
            mission_details=mission_details,
            start_date=start_date,
            start_hour=start_hour,
            end_date=end_date,
            end_hour=end_hour,
            location=location,
            facturation=facturation
        )
        
        techniciens_ids = request.POST.getlist('techniciens')
        for tech_id in techniciens_ids:
            technician = Technician.objects.get(id=tech_id)
            mission.techniciens.add(technician)
            
        # Créer une nouvelle dépense associée à la mission
        Expense.objects.create(
            mission=mission,
            hosting_days=hosting_days,
            overnight_rate=overnight_rate,
            meal_costs=meal_costs,
            transport=transport,
            shipping_costs=shipping_costs,
            various_expenses_details=various_expenses_details,
            various_expenses_price=various_expenses_price
        )
        
        return redirect('missions')


# historiques des missions validés 
class HistoryView(TemplateView):
    template_name = 'history.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        validated_missions = Mission.objects.filter(status='VALIDATED').order_by('-id')

        search_query = self.request.GET.get('search', '')
        if search_query:
            validated_missions = MissionSearchUtils.filter_missions(validated_missions, search_query)

        # Utilisation de PaginationUtils
        missions = PaginationUtils.paginate_queryset(validated_missions, self.request)

        context['missions'] = missions
        context['active_tab'] = 'history'
        return context
    
    
# Inscription
class RegisterView(TemplateView):
    template_name = "register.html"
    success_url = reverse_lazy("login")

    def post(self, request, *args, **kwargs):
        if request.method == "POST":
            username = request.POST["username"]
            first_name = request.POST["first_name"]
            last_name = request.POST["last_name"]
            poste = request.POST["poste"]
            password = request.POST["password"]
            email = request.POST["email"]

            if User.objects.filter(username=username).exists():
                messages.error(request, "Ce nom d'utilisateur existe déjà")
            else:
                user = User.objects.create_user(
                    username=username,
                    email=email, password=password,
                    first_name=first_name,
                    last_name=last_name
                    )
                
                Worker.objects.create(user_id=user, poste=poste)
                 # Si l'utilisateur est un technicien, créer aussi une entrée dans le modèle Technician
                if poste == "Techniciens":
                    Technician.objects.create(
                        first_name=first_name,
                        last_name=last_name
                    )
                messages.success(request,"Inscription réussie, connectez-vous maintenant")
                return redirect("login")

        return self.get(request, *args, **kwargs)


# Connexion
class CustomLoginView(LoginView):
    template_name = "login.html"

    def form_valid(self, form):
        response = super().form_valid(form)
        user = self.request.user

        # Vérifier si l'utilisateur a un employé associé
        return redirect("missions")


# Déconnexion
class CustomLogoutView(LogoutView):
    next_page = reverse_lazy("login")

    def dispatch(self, request, *args, **kwargs):
        messages.success(request, "Vous avez été déconnecté.")
        return super().dispatch(request, *args, **kwargs)


#  classe pour mettre à jour les données entrés par l'utilisateur 
class EditMissionView(View):
    def post(self, request, mission_id, *args, **kwargs):
        mission = Mission.objects.get(id=mission_id)
        
        # Mettre le statut à "NEW"
        mission.status = 'NEW'
        
        # Récupérer les données du formulaire
        bluedesk_link = request.POST.get('bluedesk_link')
        mission_details = request.POST.get('mission_details')
        start_date = request.POST.get('start_date')
        start_hour = request.POST.get('start_hour')
        end_date = request.POST.get('end_date')
        end_hour = request.POST.get('end_hour')
        location = request.POST.get('location')
        facturation = request.POST.get('facturation') == 'on'
        
        
        # Mettre à jour la mission
        mission.bluedesk_link = bluedesk_link
        mission.mission_details = mission_details
        mission.start_date = start_date
        mission.start_hour = start_hour
        mission.end_date = end_date
        mission.end_hour = end_hour
        mission.location = location
        mission.facturation = facturation
        mission.save()
        
        # Mise à jour des techniciens
        mission.techniciens.clear()
        techniciens_ids = request.POST.getlist('techniciens')
        for tech_id in techniciens_ids:
            technician = Technician.objects.get(id=tech_id)
            mission.techniciens.add(technician)
        
        # Mise à jour des dépenses
        try:
            expense = mission.depenses.first()  # Suppose que mission.depenses est le related_name dans le modèle Expense
            
            hosting_days = int(request.POST.get('hosting_days', 0))
            overnight_rate = Decimal(request.POST.get('overnight_rate', 0))
            meal_costs = Decimal(request.POST.get('meal_costs', 0))
            transport = request.POST.get('transport')
            shipping_costs = Decimal(request.POST.get('shipping_costs', 0))
            various_expenses_details = request.POST.get('various_expenses_details')
            various_expenses_price = Decimal(request.POST.get('various_expenses_price', 0))
            
            # Mettre à jour les valeurs
            expense.hosting_days = hosting_days
            expense.overnight_rate = overnight_rate
            expense.meal_costs = meal_costs
            expense.transport = transport
            expense.shipping_costs = shipping_costs
            expense.various_expenses_details = various_expenses_details
            expense.various_expenses_price = various_expenses_price
            
            # Recalcul des totaux (si nécessaire)
            # Note: Ceci dépend de votre modèle et de vos calculs spécifiques
            start_dt = datetime.strptime(start_date, '%Y-%m-%d')
            end_dt = datetime.strptime(end_date, '%Y-%m-%d')
            days_diff = (end_dt - start_dt).days + 1  # +1 pour inclure le jour de fin
            
            expense.total_hosting = hosting_days * overnight_rate
            expense.total_meal_costs = meal_costs * days_diff
            expense.total_expenses = (
                expense.total_hosting + 
                expense.total_meal_costs + 
                shipping_costs + 
                various_expenses_price
            )
            
            expense.save()
        except Exception as e:
            # Gérer l'erreur si nécessaire
            pass
        
        return redirect('missions')
    
    

# Pour la validation des missions
class ValidateMissionView(View):
    def post(self, request, *args, **kwargs):
        if not request.user.has_perm('app_name.can_validate_mission'):
            raise PermissionDenied
            
        mission_id = request.POST.get('mission_id')
        comment = request.POST.get('comment', '')
        mission = get_object_or_404(Mission, id=mission_id)
        
        # Mettre à jour le statut
        mission.status = 'VALIDATED'
        mission.validated_at = datetime.now()
        mission.save()
        
        # Email de notification avec adresse spécifiée
        today = datetime.now().strftime("%d/%m/%Y")
        
        # Adresse(s) email spécifiée(s) pour recevoir les notifications
        notification_email = 'mihajarazafimahazoson@gmail.com.com'  # Remplacez par l'adresse souhaitée
        
        subject = f"Demande de validation du {today}"
        message = f"La demande de mission {mission_id} du {today} a été validée par le DG."
        
        send_mail(
            subject,
            message,
            'mihajarazafimahazoson@gmail.com',  # Adresse d'expéditeur
            [notification_email],  # Liste des destinataires
            fail_silently=True,
        )
        
        messages.success(request, f"La mission a été validée avec succès.")
        return redirect(reverse('missions'))

# Pour le refus de la mission
class RefuseMissionView(View):
    def post(self, request, *args, **kwargs):
        if not request.user.has_perm('app_name.can_refuse_mission'):
            raise PermissionDenied
            
        mission_id = request.POST.get('mission_id')
        refusal_reason = request.POST.get('refusal_reason', '')
        
        if not refusal_reason.strip():
            messages.error(request, "Veuillez saisir un motif de refus.")
            return redirect(reverse('missions'))
            
        mission = get_object_or_404(Mission, id=mission_id)
        
        # Mettre à jour le statut
        mission.status = 'REFUSED'
        mission.refusal_reason = refusal_reason
        mission.refused_at = datetime.now()
        mission.save()
        
        # Email de notification avec adresse spécifiée
        today = datetime.now().strftime("%d/%m/%Y")
        
        # Adresse(s) email spécifiée(s) pour recevoir les notifications
        notification_email = 'mihajarazafimahazoson@gmail.com'  # Remplacez par l'adresse souhaitée
        
        subject = f"Demande de validation du {today}"
        message = f"La demande de mission {mission_id} du {today} a été refusée par le DG.\n\nMotif du refus : {refusal_reason}"
        
        send_mail(
            subject,
            message,
            'mihajarazafimahazoson@gmail.com',  # Adresse d'expéditeur
            [notification_email],  # Liste des destinataires
            fail_silently=True,
        )
        
        messages.success(request, f"La mission a été refusée.")
        return redirect(reverse('missions'))
    
  
#class pour le telechargement du pdf dans le modal details 
class GeneratePDFView(View):
    def get(self, request, mission_id, *args, **kwargs):
        # Récupérer la mission
        mission = get_object_or_404(Mission, id=mission_id)
        
        # Charger le template HTML
        context = {
            'mission': mission,
            'expenses': mission.depenses.all() 
        }
        
        html_string = render_to_string('pdf_template.html', context)
        
        # Créer un fichier temporaire pour stocker le PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            temp_filename = tmp.name
        
        # Générer le PDF avec WeasyPrint
        html = HTML(string=html_string)
        html.write_pdf(temp_filename)
        
        # Lire le fichier temporaire et le renvoyer dans la réponse
        with open(temp_filename, 'rb') as f:
            pdf_content = f.read()
            
        # Retourner le fichier PDF en tant que réponse HTTP
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="mission_{mission.id}.pdf"'
        
        return response
    

#class pour le telechargement du pdf de toutes les missions
class ExportMissionsPDFView(View):
    def get(self, request):
        # Récupérer toutes les missions
        missions = Mission.objects.all().prefetch_related('depenses', 'techniciens')
        
        # Préparer les données pour le résumé
        validated_count = missions.filter(status='VALIDATED').count()
        new_count = missions.filter(status='NEW').count()
        refused_count = missions.filter(status='REFUSED').count()
        
        # Calculer le total des dépenses
        total_expenses = 0
        for mission in missions:
            for expense in mission.depenses.all():
                total_expenses += expense.total_expenses
        
        # Préparer le contexte
        context = {
            'missions': missions,
            'validated_count': validated_count,
            'new_count': new_count,
            'refused_count': refused_count,
            'total_expenses': total_expenses,
            'today': date.today(),
        }
        
        # Rendre le HTML
        html_string = render_to_string('missions_pdf_export.html', context)
        
        # Générer le PDF
        html = HTML(string=html_string)
        pdf_file = html.write_pdf()
        
        # Créer la réponse HTTP
        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="missions_export.pdf"'
        
        return response
    
    
#class pour l'export Excel des missions 
class ExportMissionsExcelView(View):
    def get(self, request):
        # Créer un nouveau classeur
        wb = Workbook()
        ws = wb.active
        ws.title = "Missions"
        
        # Définir les styles
        header_font = Font(name='Arial', bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        border = Border(
            left=Side(style='thin', color='000000'),
            right=Side(style='thin', color='000000'),
            top=Side(style='thin', color='000000'),
            bottom=Side(style='thin', color='000000')
        )
        
        # En-têtes des colonnes
        headers = [
            'ID', 'Détails', 'Techniciens', 'Lieu', 'Date de début', 
            'Date de fin', 'Statut', 'Total des dépenses(Ar)'
        ]
        
        # Appliquer les en-têtes
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Définir la largeur des colonnes
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[chr(64 + col)].width = 20
        
        # Récupérer les données des missions
        missions = Mission.objects.all().prefetch_related('depenses', 'techniciens')
        
        # Ajouter les données
        row_num = 2
        for mission in missions:
            # Calculer le total des dépenses pour cette mission
            total_expenses = sum(expense.total_expenses for expense in mission.depenses.all())
            
            # Obtenir la liste des techniciens
            tech_list = ', '.join([f"{tech.first_name} {tech.last_name}" for tech in mission.techniciens.all()])
            
            # Mapper les statuts
            status_mapping = {
                'NEW': 'Nouvelle',
                'VALIDATED': 'Validée',
                'REFUSED': 'Refusée'
            }
            status_display = status_mapping.get(mission.status, mission.status)
            
            # Ajouter les données de la mission
            row = [
                mission.id,
                mission.mission_details[:100],  # Tronquer pour éviter des cellules trop longues
                tech_list,
                mission.location,
                mission.start_date.strftime('%d/%m/%Y'),
                mission.end_date.strftime('%d/%m/%Y'),
                status_display,
                f"{total_expenses:.2f} "
            ]
            
            for col_num, cell_value in enumerate(row, 1):
                cell = ws.cell(row=row_num, column=col_num)
                cell.value = cell_value
                cell.border = border
                
                # Appliquer des styles spécifiques selon le statut
                if col_num == 7:  # Colonne du statut
                    if cell_value == 'Validée':
                        cell.font = Font(color='27AE60', bold=True)
                    elif cell_value == 'Refusée':
                        cell.font = Font(color='E74C3C', bold=True)
                    elif cell_value == 'Nouvelle':
                        cell.font = Font(color='F39C12', bold=True)
            
            row_num += 1
        
        # Créer un second onglet pour les détails des dépenses
        ws_expenses = wb.create_sheet(title="Détails des dépenses")
        
        # En-têtes pour les dépenses
        expense_headers = [
            'ID Mission', 'Mission', 'Techniciens', 'Jours d\'hébergement', 'Prix nuitée', 
            'Total hébergement', 'Coût repas', 'Total repas', 'Transport', 
            'Frais de transport', 'Divers', 'Frais divers', 'Total'
        ]
        
        # Appliquer les en-têtes
        for col_num, header in enumerate(expense_headers, 1):
            cell = ws_expenses.cell(row=1, column=col_num)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Définir la largeur des colonnes
        for col in range(1, len(expense_headers) + 1):
            ws_expenses.column_dimensions[chr(64 + col)].width = 18
        
        # Ajouter les données des dépenses
        row_num = 2
        for mission in missions:
            tech_list = ', '.join([f"{tech.first_name} {tech.last_name}" for tech in mission.techniciens.all()])
            
            for expense in mission.depenses.all():
                row = [
                    mission.id,
                    mission.mission_details[:50],
                    tech_list,
                    expense.hosting_days,
                    f"{expense.overnight_rate:.2f} ",
                    f"{expense.total_hosting:.2f} ",
                    f"{expense.meal_costs:.2f} ",
                    f"{expense.total_meal_costs:.2f} ",
                    expense.transport,
                    f"{expense.shipping_costs:.2f} ",
                    expense.various_expenses_details,
                    f"{expense.various_expenses_price:.2f} ",
                    f"{expense.total_expenses:.2f} "
                ]
                
                for col_num, cell_value in enumerate(row, 1):
                    cell = ws_expenses.cell(row=row_num, column=col_num)
                    cell.value = cell_value
                    cell.border = border
                
                row_num += 1
        
        # Créer un troisième onglet pour les résumés
        ws_summary = wb.create_sheet(title="Résumé")
        
        # Styles pour les titres
        title_font = Font(name='Arial', bold=True, size=14)
        subtitle_font = Font(name='Arial', bold=True, size=12)
        
        # Titre
        ws_summary.cell(row=1, column=1).value = "Résumé des Missions"
        ws_summary.cell(row=1, column=1).font = title_font
        
        # Date de génération
        ws_summary.cell(row=2, column=1).value = f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        
        # Compteurs
        validated_count = missions.filter(status='VALIDATED').count()
        new_count = missions.filter(status='NEW').count()
        refused_count = missions.filter(status='REFUSED').count()
        
        # Total des dépenses
        total_expenses = 0
        for mission in missions:
            mission_expenses = sum(expense.total_expenses for expense in mission.depenses.all())
            total_expenses += mission_expenses
        
        # Ajouter les statistiques
        stats_row = 4
        ws_summary.cell(row=stats_row, column=1).value = "Statistiques"
        ws_summary.cell(row=stats_row, column=1).font = subtitle_font
        
        ws_summary.cell(row=stats_row+1, column=1).value = "Total des missions:"
        ws_summary.cell(row=stats_row+1, column=2).value = len(missions)
        
        ws_summary.cell(row=stats_row+2, column=1).value = "Missions validées:"
        ws_summary.cell(row=stats_row+2, column=2).value = validated_count
        
        ws_summary.cell(row=stats_row+3, column=1).value = "Nouvelles missions:"
        ws_summary.cell(row=stats_row+3, column=2).value = new_count
        
        ws_summary.cell(row=stats_row+4, column=1).value = "Missions refusées:"
        ws_summary.cell(row=stats_row+4, column=2).value = refused_count
        
        ws_summary.cell(row=stats_row+5, column=1).value = "Total des dépenses:"
        ws_summary.cell(row=stats_row+5, column=2).value = f"{total_expenses:.2f} "
        
        # Ajuster la largeur des colonnes
        for col in range(1, 3):
            ws_summary.column_dimensions[chr(64 + col)].width = 25
        
        # Sauvegarder dans un buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        # Créer la réponse HTTP
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=missions_export.xlsx'
        
        return response


#class pour l'export CSV 
class ExportMissionsCSVView(View):
    def get(self, request):
        # Créer un buffer pour écrire le CSV
        buffer = StringIO()
        writer = csv.writer(buffer)
        
        # Écrire l'en-tête
        writer.writerow([
            'ID', 'Détails de la mission', 'Techniciens', 'Lieu', 
            'Date de début', 'Date de fin', 'Statut', 'Total des dépenses(Ar)'
        ])
        
        # Récupérer toutes les missions avec leurs relations
        missions = Mission.objects.all().prefetch_related('depenses', 'techniciens')
        
        # Mapper les statuts pour l'affichage
        status_mapping = {
            'NEW': 'Nouvelle',
            'VALIDATED': 'Validée',
            'REFUSED': 'Refusée'
        }
        
        # Écrire les données des missions
        for mission in missions:
            # Calculer le total des dépenses
            total_expenses = sum(expense.total_expenses for expense in mission.depenses.all())
            
            # Obtenir la liste des techniciens formatée
            tech_list = ', '.join([f"{tech.first_name} {tech.last_name}" for tech in mission.techniciens.all()])
            
            # Convertir les dates en format lisible
            start_date = mission.start_date.strftime('%d/%m/%Y')
            end_date = mission.end_date.strftime('%d/%m/%Y')
            
            # Récupérer l'affichage du statut
            status_display = status_mapping.get(mission.status, mission.status)
            
            # Écrire la ligne
            writer.writerow([
                mission.id,
                mission.mission_details,
                tech_list,
                mission.location,
                start_date,
                end_date,
                status_display,
                f"{total_expenses:.2f}"
            ])
        
        # Préparer la réponse HTTP
        response = HttpResponse(buffer.getvalue(), content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename=missions_export_{datetime.now().strftime("%Y%m%d")}.csv'
        
        return response


#class pour l'export Word
class ExportMissionsDocxView(View):
    def get(self, request):
        # Créer un nouveau document
        doc = Document()
        
        # Configurer le style du document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Ajouter un titre
        title = doc.add_heading('Rapport des Missions', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter la date de génération
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(f'Généré le {datetime.now().strftime("%d/%m/%Y à %H:%M")}')
        
        # Ajouter un résumé
        doc.add_heading('Résumé', level=1)
        
        # Récupérer les missions
        missions = Mission.objects.all().prefetch_related('depenses', 'techniciens')
        
        # Calculer les statistiques
        total_missions = missions.count()
        validated_count = missions.filter(status='VALIDATED').count()
        new_count = missions.filter(status='NEW').count()
        refused_count = missions.filter(status='REFUSED').count()
        
        total_expenses = 0
        for mission in missions:
            mission_expenses = sum(expense.total_expenses for expense in mission.depenses.all())
            total_expenses += mission_expenses
        
        # Ajouter les statistiques
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_rows = [
            ('Total des missions:', str(total_missions)),
            ('Missions validées:', str(validated_count)),
            ('Nouvelles missions:', str(new_count)),
            ('Missions refusées:', str(refused_count)),
            ('Total des dépenses:', f"{total_expenses:.2f} Ar")
        ]
        
        for i, (label, value) in enumerate(summary_rows):
            cell = summary_table.cell(i, 0)
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
            
            cell = summary_table.cell(i, 1)
            cell.text = value
        
        # Ajouter un espace
        doc.add_paragraph()
        
        # Liste des missions
        doc.add_heading('Liste des Missions', level=1)
        
        # Mapper les statuts pour l'affichage
        status_mapping = {
            'NEW': 'Nouvelle',
            'VALIDATED': 'Validée',
            'REFUSED': 'Refusée'
        }
        
        # Parcourir chaque mission
        for mission in missions:
            # Titre de la mission
            mission_title = doc.add_heading(f'Mission #{mission.id}', level=2)
            
            # Ajouter les détails dans un tableau
            details_table = doc.add_table(rows=6, cols=2)
            details_table.style = 'Light Grid Accent 1'
            
            # Formater les dates
            start_date = mission.start_date.strftime('%d/%m/%Y')
            end_date = mission.end_date.strftime('%d/%m/%Y')
            
            # Obtenir la liste des techniciens
            tech_list = ', '.join([f"{tech.first_name} {tech.last_name}" for tech in mission.techniciens.all()])
            
            # Total des dépenses
            total_mission_expenses = sum(expense.total_expenses for expense in mission.depenses.all())
            
            # Récupérer l'affichage du statut
            status_display = status_mapping.get(mission.status, mission.status)
            
            # Remplir le tableau des détails
            details_rows = [
                ('Détails:', mission.mission_details),
                ('Lieu:', mission.location),
                ('Date de début:', start_date),
                ('Date de fin:', end_date),
                ('Techniciens:', tech_list),
                ('Statut:', status_display)
            ]
            
            for i, (label, value) in enumerate(details_rows):
                cell = details_table.cell(i, 0)
                cell.text = label
                cell.width = Inches(1.5)
                cell.paragraphs[0].runs[0].bold = True
                
                cell = details_table.cell(i, 1)
                cell.text = value
                
                # Colorer le statut
                if label == 'Statut:':
                    run = cell.paragraphs[0].runs[0]
                    if value == 'Validée':
                        run.font.color.rgb = RGBColor(39, 174, 96)  # Vert
                        run.bold = True
                    elif value == 'Refusée':
                        run.font.color.rgb = RGBColor(231, 76, 60)  # Rouge
                        run.bold = True
                    elif value == 'Nouvelle':
                        run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
                        run.bold = True
            
            # Ajouter des dépenses si elles existent
            expenses = mission.depenses.all()
            if expenses:
                doc.add_heading('Dépenses', level=3)
                
                # Créer un tableau pour les dépenses
                expense_table = doc.add_table(rows=1, cols=6)
                expense_table.style = 'Table Grid'
                
                # En-têtes
                headers = ['Nuitées', 'Hébergement', 'Repas', 'Transport', 'Divers', 'Total']
                for i, header in enumerate(headers):
                    cell = expense_table.cell(0, i)
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Ajouter chaque dépense
                for expense in expenses:
                    row_cells = expense_table.add_row().cells
                    row_cells[0].text = f"{expense.hosting_days} jours × {expense.overnight_rate:.2f} Ar"
                    row_cells[1].text = f"{expense.total_hosting:.2f} Ar"
                    row_cells[2].text = f"{expense.total_meal_costs:.2f} Ar"
                    row_cells[3].text = f"{expense.shipping_costs:.2f} Ar ({expense.transport})"
                    row_cells[4].text = f"{expense.various_expenses_price:.2f} Ar"
                    row_cells[5].text = f"{expense.total_expenses:.2f} Ar"
                    
                    # Centrer les cellules
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Ajouter une ligne pour le total de la mission
                row_cells = expense_table.add_row().cells
                row_cells[4].text = "TOTAL:"
                row_cells[4].paragraphs[0].runs[0].bold = True
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                row_cells[5].text = f"{total_mission_expenses:.2f} Ar"
                row_cells[5].paragraphs[0].runs[0].bold = True
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Ajouter un saut de page après chaque mission (sauf la dernière)
            if mission != missions.last():
                doc.add_page_break()
        
        # Sauvegarder dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Créer la réponse HTTP
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=missions_rapport.docx'
        
        return response
    
    


    
    